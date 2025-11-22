"""Core CV parsing engine."""
import hashlib
import re
from datetime import datetime, date
from typing import Optional, List, Tuple, Dict, Any
from io import BytesIO

import pdfplumber
from docx import Document
import dateparser
import phonenumbers
from email_validator import validate_email, EmailNotValidError

from app.models import (
    ParsedCandidate, ContactInfo, WorkExperience, Education,
    Skill, Certification, Language, ProficiencyLevel, EducationLevel
)


class CVParser:
    """Main CV parsing engine."""
    
    def __init__(self):
        """Initialize parser with patterns and heuristics."""
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        
        # Common section headers (very flexible patterns - match header anywhere on line)
        self.section_patterns = {
            'experience': re.compile(
                r'(?:^|\n)\s*(?:WORK\s+)?(?:EXPERIENCE|EMPLOYMENT|CAREER|PROFESSIONAL\s+(?:BACKGROUND|EXPERIENCE|HISTORY)|WORK\s*HISTORY)s?\s*:?\s*[\t ]*(?:$|\n)',
                re.IGNORECASE | re.MULTILINE
            ),
            'education': re.compile(
                r'(?:^|\n)\s*(?:EDUCATION|ACADEMIC\s*(?:BACKGROUND|HISTORY)?|QUALIFICATIONS|DEGREES?|TRAINING)\s*:?\s*[\t ]*(?:$|\n)',
                re.IGNORECASE | re.MULTILINE
            ),
            'skills': re.compile(
                r'(?:^|\n)\s*(?:(?:KEY|CORE|TECHNICAL|PROFESSIONAL)?\s*SKILLS|COMPETENC(?:IES|Y)|EXPERTISE|TECHNOLOGIES|PROFICIENCIES|ADDITIONAL\s+INFORMATION)\s*:?\s*[\t ]*(?:$|\n)',
                re.IGNORECASE | re.MULTILINE
            ),
            'certifications': re.compile(
                r'(?:^|\n)\s*(?:CERTIFICATIONS?|LICENSES?|CREDENTIALS|ACCREDITATIONS?|PROFESSIONAL\s+DEVELOPMENT)\s*:?\s*[\t ]*(?:$|\n)',
                re.IGNORECASE | re.MULTILINE
            ),
            'languages': re.compile(
                r'(?:^|\n)\s*(?:LANGUAGES?|LANGUAGE\s+(?:SKILLS|PROFICIENCY))\s*:?\s*[\t ]*(?:$|\n)',
                re.IGNORECASE | re.MULTILINE
            )
        }
        
        # Job title indicators
        self.seniority_keywords = {
            'junior': ['junior', 'jr', 'associate', 'entry', 'trainee'],
            'mid': ['developer', 'engineer', 'analyst', 'specialist', 'consultant'],
            'senior': ['senior', 'sr', 'lead', 'principal', 'staff'],
            'lead': ['head', 'director', 'vp', 'chief', 'cto', 'cio', 'manager', 'team lead']
        }
    
    async def parse_file(self, file_bytes: bytes, filename: str) -> ParsedCandidate:
        """Parse CV file and return structured candidate data."""
        from app.config import settings

        # Calculate file hash for caching
        file_hash = hashlib.sha256(file_bytes).hexdigest()

        # Extract text based on file type
        text = self._extract_text(file_bytes, filename)

        if not text or len(text.strip()) < 50:
            raise ValueError("Extracted text is too short or empty")

        # Try LLM parsing first if enabled
        if settings.LLM_ENABLED and settings.LLM_PARSE_ENABLED:
            try:
                print("🤖 Using LLM-based parsing...")
                from app.parser.llm_parser import LLMParser
                llm_parser = LLMParser()
                candidate = await llm_parser.parse_async(text, filename, file_hash)
                print("✅ LLM parsing successful")
                return candidate
            except Exception as e:
                print(f"⚠️  LLM parsing failed: {e}")
                if not settings.LLM_PARSE_FALLBACK:
                    raise
                print("🔄 Falling back to rule-based parsing...")

        # Rule-based parsing (original logic)
        print("📋 Using rule-based parsing...")
        contact = self._extract_contact_info(text)
        work = self._extract_work_experience(text)
        education = self._extract_education(text)
        skills = self._extract_skills(text)
        certs = self._extract_certifications(text)
        langs = self._extract_languages(text)

        return ParsedCandidate(
            contact=contact,
            work_experience=work,
            education=education,
            skills=skills,
            certifications=certs,
            languages=langs,
            raw_text=text,
            file_hash=file_hash,
            parsing_metadata={
                "filename": filename,
                "text_length": len(text),
                "parsed_at": datetime.utcnow().isoformat(),
                "parser": "rules"
            }
        )
    
    def _extract_text(self, file_bytes: bytes, filename: str) -> str:
        """Extract text from PDF or DOCX."""
        ext = filename.lower().split('.')[-1]

        if ext == 'pdf':
            text = self._extract_pdf(file_bytes)
        elif ext in ['docx', 'doc']:
            text = self._extract_docx(file_bytes)
        elif ext == 'txt':
            text = file_bytes.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        # DEBUG: Print extracted text to see what we're working with
        print(f"\n{'='*60}")
        print(f"EXTRACTED TEXT ({len(text)} chars)")
        print(f"{'='*60}")
        print(text[:1000])  # First 1000 chars
        print(f"{'='*60}\n")

        return text
    
    def _extract_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF using pdfplumber."""
        text_parts = []
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages[:12]:  # Respect MAX_PAGES limit
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)
    
    def _extract_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX - including tables."""
        doc = Document(BytesIO(file_bytes))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract text from tables (many CVs use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append("\t".join(row_text))

        return "\n\n".join(text_parts)
    
    def _extract_contact_info(self, text: str) -> ContactInfo:
        """Extract contact information from text."""
        # Extract name (heuristic: first line that looks like a name)
        name = self._extract_name(text)
        
        # Extract emails
        emails = self._extract_emails(text)
        
        # Extract phone numbers
        phones = self._extract_phones(text)
        
        # Extract location (heuristic: look for city, country patterns)
        location = self._extract_location(text)
        
        # Extract URLs
        linkedin = self._extract_linkedin(text)
        github = self._extract_github(text)
        portfolio = self._extract_portfolio(text)
        
        return ContactInfo(
            full_name=name,
            emails=emails,
            phones=phones,
            location=location,
            linkedin=linkedin,
            github=github,
            portfolio=portfolio
        )
    
    def _extract_name(self, text: str) -> Optional[str]:
        """Extract candidate name (first line heuristic)."""
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        if not lines:
            return None

        # First non-empty line - handle tabs and multi-column layouts
        first_line = lines[0]

        # Split on tabs first to get just the first column (usually the name)
        if '\t' in first_line:
            first_line = first_line.split('\t')[0].strip()

        # Also handle multiple spaces as column separators
        if '   ' in first_line:  # 3+ spaces
            first_line = re.split(r'\s{3,}', first_line)[0].strip()

        # Check if it looks like a name (2-4 words, capitalized)
        words = first_line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w):
            return first_line

        return None
    
    def _extract_emails(self, text: str) -> List[str]:
        """Extract email addresses."""
        emails = []
        for match in self.email_pattern.finditer(text):
            email = match.group(0)
            try:
                validated = validate_email(email, check_deliverability=False)
                emails.append(validated.normalized)
            except EmailNotValidError:
                continue
        return list(set(emails))[:3]  # Max 3 unique emails
    
    def _extract_phones(self, text: str) -> List[str]:
        """Extract phone numbers."""
        phones = []
        seen_digits = set()  # Track by raw digits to avoid duplicates

        # Try GB first for UK numbers, then other regions
        for region in ['GB', 'US', 'IN', None]:
            try:
                for match in phonenumbers.PhoneNumberMatcher(text, region):
                    # Get the raw digits to check for duplicates
                    raw_digits = ''.join(filter(str.isdigit, match.raw_string))

                    # Skip if we've already seen these digits
                    if raw_digits in seen_digits:
                        continue

                    # Validate the number is actually valid for the detected region
                    if not phonenumbers.is_valid_number(match.number):
                        continue

                    formatted = phonenumbers.format_number(
                        match.number,
                        phonenumbers.PhoneNumberFormat.INTERNATIONAL
                    )
                    phones.append(formatted)
                    seen_digits.add(raw_digits)
            except:
                continue

        return phones[:3]  # Max 3 unique phones
    
    def _extract_location(self, text: str) -> Optional[str]:
        """Extract location - only look in first few lines near contact info."""
        # Only look in first 500 chars where contact info typically is
        header_text = text[:500]

        # Common UK/US location indicators (cities, regions, countries)
        location_keywords = [
            'london', 'manchester', 'birmingham', 'leeds', 'glasgow', 'reading',
            'uk', 'united kingdom', 'england', 'scotland', 'wales',
            'usa', 'united states', 'new york', 'california', 'texas',
            'berkshire', 'warwickshire', 'surrey', 'kent'
        ]

        # Extract name words to skip
        first_line = text.split('\n')[0] if text else ''
        name_words = set()
        if first_line.strip():
            parts = re.split(r'[\t]|\s{3,}', first_line.strip())
            name_part = parts[0] if parts else ''
            for word in name_part.split():
                if word and len(word) > 1:
                    name_words.add(word.lower())

        # Pattern: "City, Region" where at least one part is a known location
        location_pattern = re.compile(
            r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'
        )

        for match in location_pattern.finditer(header_text):
            city, region = match.groups()

            # Skip if city word is in the name
            if city.split()[0].lower() in name_words:
                continue

            # At least one part must be a known location keyword
            combined = f"{city} {region}".lower()
            if any(keyword in combined for keyword in location_keywords):
                return f"{city}, {region}"

        return None
    
    def _extract_linkedin(self, text: str) -> Optional[str]:
        """Extract LinkedIn URL."""
        linkedin_pattern = re.compile(r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+/?', re.IGNORECASE)
        match = linkedin_pattern.search(text)
        if match:
            url = match.group(0)
            if not url.startswith('http'):
                url = 'https://' + url
            return url
        return None
    
    def _extract_github(self, text: str) -> Optional[str]:
        """Extract GitHub URL."""
        github_pattern = re.compile(r'(?:https?://)?(?:www\.)?github\.com/[\w-]+/?', re.IGNORECASE)
        match = github_pattern.search(text)
        if match:
            url = match.group(0)
            if not url.startswith('http'):
                url = 'https://' + url
            return url
        return None
    
    def _extract_portfolio(self, text: str) -> Optional[str]:
        """Extract portfolio/personal website URL."""
        # Generic URL pattern, excluding common social media
        url_pattern = re.compile(
            r'(?:https?://)?(?:www\.)?[\w-]+\.(?:com|io|dev|net|org)/[\w/-]*',
            re.IGNORECASE
        )
        for match in url_pattern.finditer(text[:1000]):
            url = match.group(0)
            if 'linkedin.com' not in url.lower() and 'github.com' not in url.lower():
                if not url.startswith('http'):
                    url = 'https://' + url
                return url
        return None
    
    def _extract_work_experience(self, text: str) -> List[WorkExperience]:
        """Extract work experience entries."""
        # Find experience section with more specific pattern
        exp_match = self.section_patterns['experience'].search(text)

        # DEBUG
        print(f"DEBUG: Looking for EXPERIENCE section...")
        print(f"  Pattern matched: {exp_match is not None}")
        if exp_match:
            print(f"  Match position: {exp_match.start()}-{exp_match.end()}")
            print(f"  Match text: {repr(text[exp_match.start():exp_match.end()])}")

        if not exp_match:
            return []
        
        # Get text after experience header until next major section
        exp_start = exp_match.end()
        exp_text = self._extract_section(text, exp_start)
        
        # New strategy: split by blank lines, then identify job blocks
        # A job entry typically has:
        # - Job Title (non-bullet line)
        # - Company, Location (non-bullet line)
        # - Date Range (line with Month Year - Month Year or Present)
        # - Bullets starting with •
        
        entries = []
        current_block = []
        lines = exp_text.split('\n')
        
        for line in lines:
            line = line.strip()
            
            # Blank line = end of current block
            if not line:
                if current_block:
                    entries.append('\n'.join(current_block))
                    current_block = []
                continue
            
            # Check if this looks like start of new job (has date range)
            has_date_range = bool(re.search(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\s*[-–]\s*(?:Present|Current|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)', line, re.IGNORECASE))
            
            # If we see a date range and already have content, start new block
            if has_date_range and current_block and not current_block[-1].startswith(('•', '-', '*')):
                # Save previous block
                entries.append('\n'.join(current_block))
                current_block = [line]
            else:
                current_block.append(line)
        
        # Don't forget last block
        if current_block:
            entries.append('\n'.join(current_block))
        
        # Parse each entry
        work_list = []
        for entry_text in entries:
            if len(entry_text.strip()) > 30:
                work = self._parse_work_entry(entry_text)
                if work:
                    work_list.append(work)
        
        return work_list[:10]
    
    def _extract_section(self, text: str, start_pos: int) -> str:
        """Extract text from start_pos until next major section."""
        remaining = text[start_pos:]
        
        # Find next section header
        next_section = None
        for pattern in self.section_patterns.values():
            match = pattern.search(remaining)
            if match and (next_section is None or match.start() < next_section):
                next_section = match.start()
        
        if next_section:
            return remaining[:next_section]
        return remaining
    
    def _split_experience_entries(self, text: str) -> List[str]:
        """Split experience section into individual entries."""
        # Split on double newlines or lines with company/date patterns
        entries = []
        current = []
        
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                if current:
                    entries.append('\n'.join(current))
                    current = []
            else:
                current.append(line)
        
        if current:
            entries.append('\n'.join(current))
        
        return [e for e in entries if len(e) > 20]  # Filter short fragments
    
    def _parse_work_entry(self, entry: str) -> Optional[WorkExperience]:
        """Parse a single work experience entry - flexible for multiple formats."""
        lines = [l.strip() for l in entry.split('\n') if l.strip()]
        if len(lines) < 2:
            return None

        # Strategy: Detect format dynamically
        # Format A: Title / Company / Date Range
        # Format B: Company / Title+Date / Bullets
        # Format C: Company+Date / Title / Bullets

        title = None
        company = None
        date_line = None
        bullet_start_idx = 0

        # First, identify which lines have dates
        date_pattern = re.compile(
            r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\s*[-–]\s*(?:Present|Current|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})',
            re.IGNORECASE
        )

        # Parse first few lines
        for i, line in enumerate(lines[:4]):
            # Skip bullet points
            if line.startswith(('•', '-', '*', '–')):
                bullet_start_idx = i
                break

            # Check for date range
            date_match = date_pattern.search(line)

            if date_match and not date_line:
                date_line = line
                # Extract title from line with date (before the date)
                line_before_date = line[:date_match.start()].strip()

                if i == 0:
                    # Format C: Company+Date on line 1, title on line 2
                    company = line_before_date if line_before_date else line.split()[0]
                    if len(lines) > 1:
                        title = lines[1]
                    bullet_start_idx = 2
                elif i == 1:
                    # Format B: Company on line 1, Title+Date on line 2
                    if not company:
                        company = lines[0]
                    title = line_before_date if line_before_date else "Unknown Position"
                    bullet_start_idx = 2
                else:
                    # Format A: Date on line 3+
                    if not title:
                        title = lines[0]
                    if not company:
                        company = lines[1]
                    bullet_start_idx = i + 1
                break

            elif not title:
                title = line
            elif not company:
                company = line

        # Defaults if not found
        if not title:
            title = lines[0]
        if not company:
            company = lines[1] if len(lines) > 1 else "Unknown Company"
        
        # Clean company (remove location info)
        # Handle both "Company, Location" and "Company  Location" (2+ spaces)
        if ',' in company:
            company_clean = company.split(',')[0].strip()
        elif '  ' in company:  # 2+ spaces
            company_clean = re.split(r'\s{2,}', company)[0].strip()
        else:
            company_clean = company.strip()
        
        # Extract dates from date_line or full entry
        dates = self._extract_dates(date_line if date_line else entry)
        start_date = dates[0] if dates else None
        end_date = dates[1] if len(dates) > 1 else None
        
        # Calculate duration
        duration_months = None
        if start_date:
            end = end_date if end_date else date.today()
            duration_months = (end.year - start_date.year) * 12 + (end.month - start_date.month)
            duration_months = max(1, duration_months)  # At least 1 month
        
        # Extract bullets
        bullets = []
        for l in lines[bullet_start_idx:]:
            # Clean bullet marker
            clean = re.sub(r'^[•\-*–]\s*', '', l).strip()
            # Only include actual descriptive bullets
            if clean and len(clean) > 15 and not re.match(r'^(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)', clean):
                bullets.append(clean)
        
        # Infer seniority from title
        seniority = self._infer_seniority(title)
        
        # Confidence based on what we found
        confidence = 0.5
        if dates and len(dates) >= 2:
            confidence += 0.2
        if bullets:
            confidence += 0.15
        if company_clean != "Unknown Company":
            confidence += 0.15
        
        return WorkExperience(
            employer=company_clean[:200],
            title=title[:200],
            start_date=start_date,
            end_date=end_date,
            duration_months=duration_months,
            bullets=bullets[:10],
            inferred_seniority=seniority,
            confidence=min(0.95, confidence)
        )
    
    def _extract_dates(self, text: str) -> List[date]:
        """Extract dates from text."""
        dates = []
        
        # Look for date range patterns first (most reliable)
        # Pattern: "Month Year - Month Year" or "Month Year - Present"
        range_pattern = r'(?P<start>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})\s*[-–]\s*(?P<end>(?:Present|Current|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}))'
        
        for match in re.finditer(range_pattern, text, re.IGNORECASE):
            start_str = match.group('start')
            end_str = match.group('end')
            
            try:
                start_parsed = dateparser.parse(start_str)
                if start_parsed and 1990 <= start_parsed.year <= 2030:
                    dates.append(start_parsed.date())
                
                if end_str.lower() not in ['present', 'current']:
                    end_parsed = dateparser.parse(end_str)
                    if end_parsed and 1990 <= end_parsed.year <= 2030:
                        dates.append(end_parsed.date())
                # If "Present", don't add end date (None means current)
            except:
                continue
        
        # If we found a range, return it
        if len(dates) >= 1:
            return dates[:2]
        
        # Fallback: look for individual dates
        date_patterns = [
            r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b',
            r'\b\d{1,2}/\d{4}\b',
            r'\b\d{4}\b'
        ]
        
        for pattern in date_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                try:
                    parsed = dateparser.parse(match.group(0))
                    if parsed and 1990 <= parsed.year <= 2030:
                        dates.append(parsed.date())
                except:
                    continue
        
        return sorted(set(dates))[:2]
    
    def _infer_seniority(self, title: str) -> Optional[str]:
        """Infer seniority level from job title."""
        title_lower = title.lower()
        
        for level, keywords in self.seniority_keywords.items():
            if any(kw in title_lower for kw in keywords):
                return level
        
        return 'mid'  # Default
    
    def _extract_education(self, text: str) -> List[Education]:
        """Extract education entries."""
        edu_match = self.section_patterns['education'].search(text)
        if not edu_match:
            return []
        
        edu_text = self._extract_section(text, edu_match.end())
        entries = self._split_experience_entries(edu_text)
        
        edu_list = []
        for entry in entries:
            edu = self._parse_education_entry(entry)
            if edu:
                edu_list.append(edu)
        
        return edu_list[:5]
    
    def _parse_education_entry(self, entry: str) -> Optional[Education]:
        """Parse education entry."""
        lines = [l.strip() for l in entry.split('\n') if l.strip()]
        if not lines:
            return None
        
        institution = lines[0]
        
        # Try to find degree type - use EducationLevel enum values directly
        degree_keywords = {
            EducationLevel.DOCTORATE: ['phd', 'ph.d', 'ph.d.', 'doctorate', 'doctoral', 'd.phil'],
            EducationLevel.MASTERS: ['master', "master's", 'msc', 'm.sc', 'm.sc.', 'ma', 'm.a', 'm.a.', 'mba', 'm.b.a'],
            EducationLevel.BACHELORS: ['bachelor', "bachelor's", 'bsc', 'b.sc', 'b.sc.', 'ba', 'b.a', 'b.a.', 'bs', 'b.s', 'b.s.'],
            EducationLevel.ASSOCIATES: ['associate', "associate's", 'as', 'a.s', 'a.s.'],
        }
        
        degree = None
        entry_lower = entry.lower()
        for deg_type, keywords in degree_keywords.items():
            if any(kw in entry_lower for kw in keywords):
                degree = deg_type
                break
        
        # Extract dates
        dates = self._extract_dates(entry)
        
        return Education(
            institution=institution[:200],
            degree=degree,
            field=None,  # TODO: extract field
            start_date=dates[0] if dates else None,
            end_date=dates[1] if len(dates) > 1 else None,
            confidence=0.6
        )
    
    def _extract_skills(self, text: str) -> List[Skill]:
        """Extract skills (improved version with better tech skill detection)."""
        skills_match = self.section_patterns['skills'].search(text)
        
        if skills_match:
            skills_text = self._extract_section(text, skills_match.end())
        else:
            # Fallback: use full text
            skills_text = text
        
        # Expanded tech skills dictionary with categories
        tech_skills_db = {
            # Languages
            'python': ['python', 'python3', 'py'],
            'javascript': ['javascript', 'js', 'ecmascript'],
            'typescript': ['typescript', 'ts'],
            'java': ['java'],
            'csharp': ['c#', 'csharp', 'c-sharp'],
            'cpp': ['c++', 'cpp'],
            'go': ['golang', 'go'],
            'rust': ['rust'],
            'sql': ['sql', 'mysql', 'postgresql', 'postgres'],
            
            # Frameworks
            'react': ['react', 'reactjs', 'react.js'],
            'angular': ['angular', 'angularjs'],
            'vue': ['vue', 'vuejs', 'vue.js'],
            'django': ['django'],
            'flask': ['flask'],
            'fastapi': ['fastapi', 'fast api'],
            'nodejs': ['node', 'nodejs', 'node.js'],
            'express': ['express', 'expressjs'],
            
            # Cloud & DevOps
            'aws': ['aws', 'amazon web services'],
            'azure': ['azure', 'microsoft azure'],
            'gcp': ['gcp', 'google cloud', 'google cloud platform'],
            'docker': ['docker', 'containerization'],
            'kubernetes': ['kubernetes', 'k8s'],
            'terraform': ['terraform'],
            'jenkins': ['jenkins'],
            'cicd': ['ci/cd', 'ci-cd', 'continuous integration'],
            
            # Databases
            'postgresql': ['postgresql', 'postgres', 'psql'],
            'mongodb': ['mongodb', 'mongo'],
            'redis': ['redis'],
            'elasticsearch': ['elasticsearch', 'elastic'],
            
            # Tools & Methodologies
            'git': ['git', 'github', 'gitlab'],
            'linux': ['linux', 'unix'],
            'agile': ['agile', 'scrum', 'kanban'],
        }
        
        found_skills = []
        skills_text_lower = skills_text.lower()
        
        for canonical, variations in tech_skills_db.items():
            for variant in variations:
                # Use word boundaries for better matching
                pattern = r'\b' + re.escape(variant) + r'\b'
                if re.search(pattern, skills_text_lower, re.IGNORECASE):
                    found_skills.append(Skill(
                        name=variant.title() if len(variant) > 3 else variant.upper(),
                        canonical_id=canonical,
                        group='technical',
                        confidence=0.85
                    ))
                    break  # Don't double-count variants
        
        return found_skills[:30]
    
    def _extract_skills_fuzzy(self, text: str) -> List[Skill]:
        """Fuzzy skill extraction from full text."""
        # Simplified version - just look for tech keywords
        return []
    
    def _extract_certifications(self, text: str) -> List[Certification]:
        """Extract certifications."""
        cert_match = self.section_patterns['certifications'].search(text)
        if not cert_match:
            return []
        
        cert_text = self._extract_section(text, cert_match.end())
        
        # Common certifications
        known_certs = ['aws', 'azure', 'gcp', 'pmp', 'scrum', 'cissp', 'cisa']
        
        certs = []
        for line in cert_text.split('\n'):
            line = line.strip()
            if len(line) > 5:
                cert_name = line
                for known in known_certs:
                    if known in line.lower():
                        cert_name = known.upper()
                        break
                
                dates = self._extract_dates(line)
                certs.append(Certification(
                    name=cert_name[:200],
                    issue_date=dates[0] if dates else None,
                    confidence=0.5
                ))
        
        return certs[:10]
    
    def _extract_languages(self, text: str) -> List[Language]:
        """Extract language proficiency - FIXED with better parsing."""
        lang_match = self.section_patterns['languages'].search(text)
        if not lang_match:
            return []
        
        lang_text = self._extract_section(text, lang_match.end())
        
        proficiency_map = {
            'native': ProficiencyLevel.NATIVE,
            'fluent': ProficiencyLevel.FLUENT,
            'professional': ProficiencyLevel.PROFESSIONAL,
            'intermediate': ProficiencyLevel.INTERMEDIATE,
            'basic': ProficiencyLevel.BASIC,
            'conversational': ProficiencyLevel.INTERMEDIATE,
            'working': ProficiencyLevel.PROFESSIONAL,
        }
        
        languages = []
        for line in lang_text.split('\n'):
            line = line.strip()
            if len(line) < 3:
                continue
            
            # Skip if it's another section header
            if any(pattern.match(line) for pattern in self.section_patterns.values()):
                break
            
            # Try to find proficiency level
            prof = ProficiencyLevel.INTERMEDIATE  # Default
            line_lower = line.lower()
            for keyword, level in proficiency_map.items():
                if keyword in line_lower:
                    prof = level
                    break
            
            # Extract language name (before dash or colon)
            if ' - ' in line:
                lang_name = line.split(' - ')[0].strip()
            elif ':' in line:
                lang_name = line.split(':')[0].strip()
            else:
                lang_name = line.split()[0] if line.split() else line
            
            if lang_name and len(lang_name) > 1:
                languages.append(Language(
                    name=lang_name.title(),
                    proficiency=prof,
                    confidence=0.7
                ))
        
        return languages[:10]